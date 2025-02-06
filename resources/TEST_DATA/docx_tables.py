from docx import Document
import pandas as pd

# DOCX_FILE = "p9467_doc2__9-7-17__Antigua_and_Barbuda_CD_9467_ProDoc__revised.docx"
DOCX_FILE = "p9467_doc0__4-1-16__Antigua_Barbuda_CD_9467_PIF.docx"


def main():
    doc = Document(DOCX_FILE)

    tables = doc.tables

    print(f"Number of tables: {len(tables)}")
    print(tables)


def extract_docx_tables(docx_path):
    doc = Document(docx_path)
    tables = []

    for table in doc.tables:
        data = []

        # Get headers
        headers = []
        for cell in table.rows[0].cells:
            headers.append(cell.text.strip())

        # Get data rows
        for row in table.rows[1:]:
            row_data = []
            for cell in row.cells:
                # Handle nested tables if they exist
                if cell.tables:
                    nested_text = []
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            nested_text.append(
                                " ".join(c.text.strip() for c in nested_row.cells)
                            )
                    row_data.append("\n".join(nested_text))
                else:
                    row_data.append(cell.text.strip())
            data.append(row_data)

        # Create DataFrame
        df = pd.DataFrame(data, columns=headers)
        tables.append(df)

    return tables


def save_tables_csv(tables: list[pd.DataFrame], output_dir: str = "outputs"):
    for i, table in enumerate(tables):
        table.to_csv(f"{output_dir}/table_{i + 1}.csv", index=False)


if __name__ == "__main__":
    tables = extract_docx_tables(DOCX_FILE)
    # for i, table in enumerate(tables):
    #     print(f"Table {i + 1}:")
    #     print(table)
    #     print("\n")
    save_tables_csv(tables)
